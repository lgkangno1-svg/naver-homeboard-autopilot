# -*- coding: utf-8 -*-
"""docx 샘플 구조 파악용 스크립트"""
import glob, sys
from docx import Document

sys.stdout.reconfigure(encoding="utf-8")

files = sorted(glob.glob(r"C:\Users\tnfwo\Desktop\홈판분석\*.docx"))
print(f"파일 수: {len(files)}")
for f in files:
    print("=" * 70)
    print("FILE:", f.split("\\")[-1])
    try:
        doc = Document(f)
    except Exception as e:
        print("  열기 실패:", e)
        continue
    paras = [p.text.strip() for p in doc.paragraphs]
    nonempty = [p for p in paras if p]
    print(f"  문단 수: {len(paras)} / 빈 줄 제외: {len(nonempty)} / 인라인 이미지: {len(doc.inline_shapes)} / 표: {len(doc.tables)}")
    print("  --- 앞 25개 문단 ---")
    shown = 0
    for p in nonempty:
        line = p.replace("\n", " ⏎ ")[:90]
        print("   |", line)
        shown += 1
        if shown >= 25:
            break
